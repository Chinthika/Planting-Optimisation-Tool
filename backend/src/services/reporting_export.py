"""
Reporting export service — generates DOCX and PDF formatted reports
from a FarmReportContract.
"""

import io

from docx import Document
from docx.shared import Pt
from fpdf import FPDF, XPos, YPos

from src.domains.reporting import FarmReportContract


def generate_docx(report: FarmReportContract) -> bytes:
    """Generates a DOCX formatted report from a FarmReportContract.
    Returns the file as bytes.
    """
    doc = Document()

    # Title
    doc.add_heading("Planting Optimisation Tool - Farm Report", level=1)
    doc.add_paragraph(f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    # Farm details section
    doc.add_heading("Farm Profile", level=2)
    farm = report.farm
    farm_details = [
        ("Farm ID", str(farm.id)),
        ("Rainfall", f"{farm.rainfall_mm} mm"),
        ("Temperature", f"{farm.temperature_celsius} °C"),
        ("Elevation", f"{farm.elevation_m} m"),
        ("pH", str(farm.ph)),
        ("Soil Texture", farm.soil_texture),
        ("Area", f"{farm.area_ha} ha"),
        ("Latitude", str(farm.latitude)),
        ("Longitude", str(farm.longitude)),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Property"
    header_cells[1].text = "Value"

    for label, value in farm_details:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    doc.add_paragraph("")

    # Recommendations section
    doc.add_heading("Species Recommendations", level=2)

    if not report.recommendations:
        doc.add_paragraph("No recommendations available for this farm.")
    else:
        rec_table = doc.add_table(rows=1, cols=4)
        rec_table.style = "Table Grid"
        header_cells = rec_table.rows[0].cells
        header_cells[0].text = "Rank"
        header_cells[1].text = "Species"
        header_cells[2].text = "Common Name"
        header_cells[3].text = "Score"

        for rec in report.recommendations:
            row_cells = rec_table.add_row().cells
            row_cells[0].text = str(rec.rank_overall)
            row_cells[1].text = rec.species_name
            row_cells[2].text = rec.species_common_name
            row_cells[3].text = f"{rec.score_mcda:.2f}"

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_pdf(report: FarmReportContract) -> bytes:
    """Generates a PDF formatted report from a FarmReportContract.
    Returns the file as bytes.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(0, 10, "Planting Optimisation Tool - Farm Report", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 8, f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M UTC')}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(5)

    # Farm details section
    pdf.set_font("Helvetica", style="B", size=13)
    pdf.cell(0, 10, "Farm Profile", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    farm = report.farm
    farm_details = [
        ("Farm ID", str(farm.id)),
        ("Rainfall", f"{farm.rainfall_mm} mm"),
        ("Temperature", f"{farm.temperature_celsius} C"),
        ("Elevation", f"{farm.elevation_m} m"),
        ("pH", str(farm.ph)),
        ("Soil Texture", farm.soil_texture),
        ("Area", f"{farm.area_ha} ha"),
        ("Latitude", str(farm.latitude)),
        ("Longitude", str(farm.longitude)),
    ]

    pdf.set_font("Helvetica", size=10)
    col_w = 60
    for label, value in farm_details:
        pdf.set_font("Helvetica", style="B", size=10)
        pdf.cell(col_w, 8, label, border=1)
        pdf.set_font("Helvetica", size=10)
        pdf.cell(0, 8, value, border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.ln(5)

    # Recommendations section
    pdf.set_font("Helvetica", style="B", size=13)
    pdf.cell(0, 10, "Species Recommendations", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if not report.recommendations:
        pdf.set_font("Helvetica", size=10)
        pdf.cell(0, 8, "No recommendations available for this farm.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    else:
        # Table headers
        pdf.set_font("Helvetica", style="B", size=10)
        pdf.cell(15, 8, "Rank", border=1)
        pdf.cell(60, 8, "Species", border=1)
        pdf.cell(60, 8, "Common Name", border=1)
        pdf.cell(0, 8, "Score", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font("Helvetica", size=10)
        for rec in report.recommendations:
            pdf.cell(15, 8, str(rec.rank_overall), border=1)
            pdf.cell(60, 8, rec.species_name[:30], border=1)
            pdf.cell(60, 8, rec.species_common_name[:30], border=1)
            pdf.cell(0, 8, f"{rec.score_mcda:.2f}", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    return bytes(pdf.output())


def generate_all_farms_docx(reports: list[FarmReportContract]) -> bytes:
    """Generates a single DOCX report covering all farms, one section per farm.
    Returns the file as bytes.
    """
    doc = Document()

    # Document title
    doc.add_heading("Planting Optimisation Tool - All Farms Report", level=1)
    if reports:
        doc.add_paragraph(f"Generated: {reports[0].generated_at.strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(f"Total farms: {len(reports)}")
    doc.add_paragraph("")

    for report in reports:
        farm = report.farm

        # Farm section heading
        doc.add_heading(f"Farm {farm.id}", level=2)

        # Farm profile table
        doc.add_heading("Farm Profile", level=3)
        farm_details = [
            ("Farm ID", str(farm.id)),
            ("Rainfall", f"{farm.rainfall_mm} mm"),
            ("Temperature", f"{farm.temperature_celsius} °C"),
            ("Elevation", f"{farm.elevation_m} m"),
            ("pH", str(farm.ph)),
            ("Soil Texture", farm.soil_texture),
            ("Area", f"{farm.area_ha} ha"),
            ("Latitude", str(farm.latitude)),
            ("Longitude", str(farm.longitude)),
        ]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Property"
        header_cells[1].text = "Value"

        for label, value in farm_details:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_paragraph("")

        # Recommendations table
        doc.add_heading("Species Recommendations", level=3)

        if not report.recommendations:
            doc.add_paragraph("No recommendations available for this farm.")
        else:
            rec_table = doc.add_table(rows=1, cols=4)
            rec_table.style = "Table Grid"
            rec_header = rec_table.rows[0].cells
            rec_header[0].text = "Rank"
            rec_header[1].text = "Species"
            rec_header[2].text = "Common Name"
            rec_header[3].text = "Score"

            for rec in report.recommendations:
                row_cells = rec_table.add_row().cells
                row_cells[0].text = str(rec.rank_overall)
                row_cells[1].text = rec.species_name
                row_cells[2].text = rec.species_common_name
                row_cells[3].text = f"{rec.score_mcda:.2f}"

        # Page break between farms (except after the last one)
        if report != reports[-1]:
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_all_farms_pdf(reports: list[FarmReportContract]) -> bytes:
    """Generates a single PDF report covering all farms, one page per farm.
    Returns the file as bytes.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    for index, report in enumerate(reports):
        pdf.add_page()
        farm = report.farm

        # Farm section title
        pdf.set_font("Helvetica", style="B", size=16)
        pdf.cell(0, 10, f"Farm {farm.id} - Planting Report", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", size=10)
        pdf.cell(0, 8, f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M UTC')}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 8, f"Farm {index + 1} of {len(reports)}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(5)

        # Farm profile section
        pdf.set_font("Helvetica", style="B", size=13)
        pdf.cell(0, 10, "Farm Profile", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        farm_details = [
            ("Farm ID", str(farm.id)),
            ("Rainfall", f"{farm.rainfall_mm} mm"),
            ("Temperature", f"{farm.temperature_celsius} C"),
            ("Elevation", f"{farm.elevation_m} m"),
            ("pH", str(farm.ph)),
            ("Soil Texture", farm.soil_texture),
            ("Area", f"{farm.area_ha} ha"),
            ("Latitude", str(farm.latitude)),
            ("Longitude", str(farm.longitude)),
        ]

        col_w = 60
        for label, value in farm_details:
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.cell(col_w, 8, label, border=1)
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 8, value, border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.ln(5)

        # Recommendations section
        pdf.set_font("Helvetica", style="B", size=13)
        pdf.cell(0, 10, "Species Recommendations", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        if not report.recommendations:
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 8, "No recommendations available for this farm.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            # Table headers
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.cell(15, 8, "Rank", border=1)
            pdf.cell(60, 8, "Species", border=1)
            pdf.cell(60, 8, "Common Name", border=1)
            pdf.cell(0, 8, "Score", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            pdf.set_font("Helvetica", size=10)
            for rec in report.recommendations:
                pdf.cell(15, 8, str(rec.rank_overall), border=1)
                pdf.cell(60, 8, rec.species_name[:30], border=1)
                pdf.cell(60, 8, rec.species_common_name[:30], border=1)
                pdf.cell(0, 8, f"{rec.score_mcda:.2f}", border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    return bytes(pdf.output())
